"""Export agent results to Markdown, YAML metadata, and DOCX."""

import sys
import yaml
from datetime import datetime
from pathlib import Path


# ── Stdout log capture ──

class _TeeWriter:
    """Write to both the original stdout and a log file."""
    def __init__(self, log_path, orig_stdout):
        self._log = open(log_path, "w", encoding="utf-8")
        self._orig = orig_stdout

    def write(self, s):
        self._orig.write(s)
        self._log.write(s)

    def flush(self):
        self._orig.flush()
        self._log.flush()

    def close(self):
        self._log.close()


_tee = None
_orig_stdout = None


def start_log(output_dir: Path):
    """Start capturing stdout to output_dir/log.txt."""
    global _tee, _orig_stdout
    _orig_stdout = sys.stdout
    _tee = _TeeWriter(output_dir / "log.txt", _orig_stdout)
    sys.stdout = _tee


def stop_log():
    """Stop capturing and restore original stdout."""
    global _tee, _orig_stdout
    if _tee is not None:
        sys.stdout = _orig_stdout
        _tee.close()
        _tee = None
        _orig_stdout = None


def _add_formatted_runs(paragraph, text):
    """Parse inline **bold** and *italic* markdown into docx runs."""
    import re
    # Split on **bold** and *italic* markers
    parts = re.split(r'(\*\*.*?\*\*|\*.*?\*)', text)
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith("*") and part.endswith("*"):
            run = paragraph.add_run(part[1:-1])
            run.italic = True
        else:
            paragraph.add_run(part)


def _md_lines_to_docx(doc, text):
    """Convert simple Markdown text to docx paragraphs."""
    for line in text.split("\n"):
        stripped = line.strip()
        if stripped.startswith("### "):
            doc.add_heading(stripped[4:], level=3)
        elif stripped.startswith("## "):
            doc.add_heading(stripped[3:], level=2)
        elif stripped.startswith("# "):
            doc.add_heading(stripped[2:], level=1)
        elif stripped.startswith("- ") or stripped.startswith("* "):
            p = doc.add_paragraph(style="List Bullet")
            _add_formatted_runs(p, stripped[2:])
        elif stripped:
            p = doc.add_paragraph()
            _add_formatted_runs(p, stripped)


def save_all(result: dict, config: dict, output_dir: Path, elapsed: float):
    """Save Markdown, meta.yaml, and DOCX report to *output_dir*."""
    ts = output_dir.name
    agents_str = ", ".join(
        f'{k}={v["model"]}' for k, v in config["agents"].items()
    )

    # ── Markdown ──
    md_path = output_dir / f"report_{ts}.md"
    with open(md_path, "w", encoding="utf-8") as f:
        f.write("# MBA Strategy Report v4\n\n")
        f.write(f"**Generated:** {datetime.now().isoformat()}\n")
        f.write(f"**Agents:** {agents_str}\n\n---\n\n")
        f.write(result.get("recommendation", ""))
        f.write("\n\n---\n\n")
        f.write(result.get("action_plan", ""))
    print(f"Markdown: {md_path}")

    # ── meta.yaml ──
    with open(output_dir / "meta.yaml", "w", encoding="utf-8") as f:
        yaml.dump(
            {
                "timestamp": datetime.now().isoformat(),
                "input_query": config["input_query"],
                "auto_approve": False,
                "models": {k: v["model"] for k, v in config["agents"].items()},
                "elapsed_seconds": round(elapsed, 1),
            },
            f,
            default_flow_style=False,
            sort_keys=False,
        )
    print(f"Meta:     {output_dir / 'meta.yaml'}")

    # ── DOCX (optional) ──
    try:
        from docx import Document as DocxDocument
        from docx.shared import Pt, Cm, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        doc = DocxDocument()
        for section in doc.sections:
            section.top_margin = Cm(2.5)
            section.bottom_margin = Cm(2.5)
            section.left_margin = Cm(2.5)
            section.right_margin = Cm(2.5)

        style = doc.styles["Normal"]
        style.font.name = "Calibri"
        style.font.size = Pt(11)

        def _cover_para(text, size=11, color=RGBColor(0, 0, 0),
                        bold=False, italic=False,
                        space_before=None, space_after=None):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            if space_before is not None:
                p.paragraph_format.space_before = space_before
            if space_after is not None:
                p.paragraph_format.space_after = space_after
            run = p.add_run(text)
            run.font.size = Pt(size)
            run.font.color.rgb = color
            run.bold = bold
            run.italic = italic
            return p

        # — Cover page —
        NAVY = RGBColor(31, 78, 121)
        ACCENT = RGBColor(47, 117, 181)
        DARK_GRAY = RGBColor(89, 89, 89)
        LIGHT_GRAY = RGBColor(140, 140, 140)

        # Title pushed down ~40% with space_before
        _cover_para("MBA STRATEGY REPORT", size=36, color=NAVY,
                     bold=True, space_before=Cm(6), space_after=Pt(4))

        # Thin divider line
        _cover_para("━" * 30, size=10, color=ACCENT,
                     space_before=Pt(0), space_after=Pt(12))

        # Business question
        _cover_para(f"\u201c{config['input_query']}\u201d",
                     size=14, color=DARK_GRAY, italic=True,
                     space_after=Cm(1.5))

        # Date
        _cover_para(
            f"Generated: {datetime.now().strftime('%B %d, %Y')}",
            size=10, color=LIGHT_GRAY, space_after=Pt(2))

        # Models
        _cover_para(f"Models: {agents_str}",
                     size=10, color=LIGHT_GRAY, space_after=Pt(0))

        doc.add_page_break()

        recommendation = result.get("recommendation", "")
        if recommendation:
            _md_lines_to_docx(doc, recommendation)

        doc.add_page_break()

        action_plan = result.get("action_plan", "")
        if action_plan:
            _md_lines_to_docx(doc, action_plan)

        docx_path = output_dir / f"report_{ts}.docx"
        doc.save(str(docx_path))
        print(f"DOCX:     {docx_path}")
    except ImportError:
        print("[SKIP] python-docx not installed")
    except Exception as e:
        print(f"[WARNING] DOCX generation failed: {e}")

    print(f"Log:      {output_dir / 'log.txt'}")
    print(f"\nAll outputs saved to: {output_dir}")
